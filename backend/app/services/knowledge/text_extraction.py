"""Plain-text extraction for knowledge-document source files.

Mirrors the file-loading half of noledge's ``src/lib/ingest/ingest.ts``: turn an
uploaded ``txt`` / ``md`` / ``pdf`` / ``docx`` / ``csv`` blob into the plain UTF-8
text that the chunker + embedder consume. Each format uses a maintained parser
(``pypdf`` for PDF, ``python-docx`` for DOCX, the stdlib ``csv`` module for CSV);
``txt`` / ``md`` are decoded directly.

Extraction never raises on a malformed document body: parser failures are logged
and surfaced as :class:`TextExtractionError` so the calling ingest flow can roll
back cleanly instead of leaking a third-party stack trace.
"""

from __future__ import annotations

import csv
import io
from pathlib import PurePosixPath

import structlog

logger = structlog.get_logger()

# Extensions we can extract. Markdown is treated as plain text (its structure is
# meaningful to the chunker's paragraph/line separators, so we keep it verbatim).
SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({"txt", "md", "markdown", "pdf", "docx", "csv"})

# MIME → canonical extension, for callers that only have a content type.
_CONTENT_TYPE_TO_EXT: dict[str, str] = {
    "text/plain": "txt",
    "text/markdown": "md",
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/csv": "csv",
    "application/csv": "csv",
}


class TextExtractionError(Exception):
    """Raised when a document body cannot be parsed into text."""


def _decode_text(data: bytes) -> str:
    """Decode bytes as UTF-8, replacing undecodable bytes rather than failing."""
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    """Extract text from a PDF using ``pypdf`` (maintained, pure-Python)."""
    from pypdf import PdfReader
    from pypdf.errors import PyPdfError

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except (PyPdfError, ValueError, OSError) as exc:
        raise TextExtractionError("Could not read PDF document.") from exc
    return "\n\n".join(page.strip() for page in pages if page.strip())


def _extract_docx(data: bytes) -> str:
    """Extract paragraph + table text from a DOCX using ``python-docx``."""
    import zipfile

    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(io.BytesIO(data))
    except (PackageNotFoundError, zipfile.BadZipFile, ValueError, KeyError, OSError) as exc:
        raise TextExtractionError("Could not read DOCX document.") from exc

    parts: list[str] = [para.text for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                parts.append(line)
    return "\n\n".join(parts)


def _extract_csv(data: bytes) -> str:
    """Flatten CSV rows into newline-delimited, comma-joined text."""
    text = _decode_text(data)
    try:
        rows = list(csv.reader(io.StringIO(text)))
    except csv.Error as exc:
        raise TextExtractionError("Could not parse CSV document.") from exc
    lines = [", ".join(cell.strip() for cell in row) for row in rows]
    return "\n".join(line for line in lines if line.strip())


def normalize_extension(*, filename: str | None, content_type: str | None) -> str | None:
    """Resolve a canonical extension from a filename and/or MIME type."""
    if filename:
        suffix = PurePosixPath(filename).suffix.lower().lstrip(".")
        if suffix in SUPPORTED_EXTENSIONS:
            return "md" if suffix == "markdown" else suffix
    if content_type:
        base = content_type.split(";", 1)[0].strip().lower()
        mapped = _CONTENT_TYPE_TO_EXT.get(base)
        if mapped:
            return mapped
    return None


def extract_text(
    data: bytes,
    *,
    filename: str | None = None,
    content_type: str | None = None,
) -> str:
    """Extract plain text from a document blob.

    Args:
        data: Raw file bytes.
        filename: Original filename, used to infer the format by extension.
        content_type: MIME type, used as a fallback when no usable filename.

    Returns:
        Extracted, stripped UTF-8 text.

    Raises:
        TextExtractionError: The format is unsupported or the body cannot be parsed.
    """
    ext = normalize_extension(filename=filename, content_type=content_type)
    if ext is None:
        raise TextExtractionError(
            "Unsupported document type; expected one of txt, md, pdf, docx, csv."
        )

    if ext in ("txt", "md"):
        return _decode_text(data).strip()
    if ext == "pdf":
        return _extract_pdf(data).strip()
    if ext == "docx":
        return _extract_docx(data).strip()
    if ext == "csv":
        return _extract_csv(data).strip()

    # Unreachable: ext is constrained to SUPPORTED_EXTENSIONS above.
    raise TextExtractionError(f"Unsupported document type: {ext}")  # pragma: no cover
